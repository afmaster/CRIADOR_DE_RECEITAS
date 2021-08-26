import os

from kivy.app import App
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.lang import Builder
from kivy.clock import Clock

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from docx.shared import Pt



class Pannel(Screen):
    def start_func(self, paciente, medicamento, quantidade, posologia, obs, obs_2, chk):
        pic = "ineuro.jpg"
        comprador = 'comprador.jpg'
        fornecedor = 'fornecedor.jpg'

        if chk ==True:
            agora = datetime.now().strftime("%d/%m/%Y")
            agora = str(agora)
        else:
            agora = ""

        #paciente = "NOME COMPLETO DO PACIENTE"
        #medicamento = "Tramadol, 50 mg"
        #quantidade = "1 cx"
        #posologia = "1"
        #obs = "8/8 hs, se dor forte."
        #obs_2 = ""

        document = Document()

        section = document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        section = document.sections[0]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')

        ########################################3

        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells_row_0 = table.rows[0].cells
        paragraph_table = hdr_cells_row_0[0].paragraphs[0]
        run = paragraph_table.add_run()
        run.add_picture(pic, width = 2200000, height = 700000)

        hdr_cells_row_0[1].text = """
        
        Receituário de Controle Especial
        1a Via - Farmácia.
        """
        paragraph_controle =hdr_cells_row_0[1].paragraphs[0]
        run = paragraph_controle.runs
        font = run[0].font
        font.size = Pt(8)

        ######################################################

        table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        hdr_cells_row_0 = table.rows[0].cells
        hdr_cells_row_0[0].text = """
        CNPJ 13.122.535/0001-87 - CREMERS 5503
        Rua Dr. Luis Bastos do Prado, nº 1586/B, 5º andar
        Gravataí - CEP:94010-020
        Telefone (51) 3484-1745  (51) 34841756
        """
        paragraph_cnpj =hdr_cells_row_0[0].paragraphs[0]
        run = paragraph_cnpj.runs
        font = run[0].font
        font.size = Pt(8)

        table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        hdr_cells_row_0 = table.rows[0].cells
        hdr_cells_row_0[0].text = """
        PACIENTE: %s
        
        Uso Interno
        1- %s ------------------------ %s
        Tomar %s comp, VO, por dia.
        %s
        %s
        
        
        
        
        %s
        """ %(paciente,medicamento, quantidade, posologia, obs, obs_2, agora)

        #################################################
        document.add_paragraph('')

        table_8 = document.add_table(rows=1, cols=2)
        table_8.style = 'Table Grid'
        hdr_cells_row_0_8 = table_8.rows[0].cells
        paragraph_table_8 = hdr_cells_row_0_8[0].paragraphs[0]
        run = paragraph_table_8.add_run()
        run.add_picture(comprador, width=1800000, height=1300000)

        hdr_cells_row_0_8_1 = table_8.rows[0].cells
        paragraph_table_8_1 = hdr_cells_row_0_8_1[1].paragraphs[0]
        run = paragraph_table_8_1.add_run()
        run.add_picture(fornecedor, width=1800000, height=1300000)

        document.add_paragraph('')

        ################################## segunda via ######################33
        ########################################3

        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells_row_0 = table.rows[0].cells
        paragraph_table = hdr_cells_row_0[0].paragraphs[0]
        run = paragraph_table.add_run()
        run.add_picture(pic, width = 2200000, height = 700000)

        hdr_cells_row_0[1].text = """
        
        Receituário de Controle Especial
        2a Via - Paciente.
        """
        paragraph_controle =hdr_cells_row_0[1].paragraphs[0]
        run = paragraph_controle.runs
        font = run[0].font
        font.size = Pt(8)

        ######################################################

        table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        hdr_cells_row_0 = table.rows[0].cells
        hdr_cells_row_0[0].text = """
        CNPJ 13.122.535/0001-87 - CREMERS 5503
        Rua Dr. Luis Bastos do Prado, nº 1586/B, 5º andar
        Gravataí - CEP:94010-020
        Telefone (51) 3484-1745  (51) 34841756
        """
        paragraph_cnpj =hdr_cells_row_0[0].paragraphs[0]
        run = paragraph_cnpj.runs
        font = run[0].font
        font.size = Pt(8)

        table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        hdr_cells_row_0 = table.rows[0].cells
        hdr_cells_row_0[0].text = """
        PACIENTE: %s
        
        Uso Interno
        1- %s ------------------------ %s
        Tomar %s comp, VO, por dia.
        %s
        %s
        
        
        
        
        %s
        """ %(paciente,medicamento, quantidade, posologia, obs, obs_2, agora)

        #################################################
        document.add_paragraph('')


        table_8 = document.add_table(rows=1, cols=2)
        table_8.style = 'Table Grid'
        hdr_cells_row_0_8 = table_8.rows[0].cells
        paragraph_table_8 = hdr_cells_row_0_8[0].paragraphs[0]
        run = paragraph_table_8.add_run()
        run.add_picture(comprador, width=1800000, height=1300000)

        hdr_cells_row_0_8_1 = table_8.rows[0].cells
        paragraph_table_8_1 = hdr_cells_row_0_8_1[1].paragraphs[0]
        run = paragraph_table_8_1.add_run()
        run.add_picture(fornecedor, width=1800000, height=1300000)



        document.save('prescription.docx')

    def printing (self):
        os.startfile('prescription.docx', 'print')

    def opening (self):
        try:
            os.startfile('prescription.docx')
        except Exception as err:
            print(err)


class WindowManager(ScreenManager):
    pass


class Main(App):
    def build(self):
        pass

if __name__=='__main__':
    Main().run()